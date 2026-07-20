"""Rule-based CV text extraction and skill/title parsing (no AI)."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

# Common tech skills / tools (lowercase keys, display labels)
SKILL_CATALOG: dict[str, str] = {
    # Languages
    "python": "Python",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "java": "Java",
    "kotlin": "Kotlin",
    "swift": "Swift",
    "go": "Go",
    "golang": "Go",
    "rust": "Rust",
    "c++": "C++",
    "c#": "C#",
    "csharp": "C#",
    "ruby": "Ruby",
    "php": "PHP",
    "scala": "Scala",
    "matlab": "MATLAB",
    "sql": "SQL",
    "bash": "Bash",
    "shell": "Shell",
    # Web / frontend
    "react": "React",
    "react.js": "React",
    "reactjs": "React",
    "vue": "Vue",
    "vue.js": "Vue",
    "angular": "Angular",
    "next.js": "Next.js",
    "nextjs": "Next.js",
    "svelte": "Svelte",
    "html": "HTML",
    "css": "CSS",
    "tailwind": "Tailwind",
    "sass": "Sass",
    "redux": "Redux",
    # Backend / frameworks
    "fastapi": "FastAPI",
    "django": "Django",
    "flask": "Flask",
    "express": "Express",
    "express.js": "Express",
    "node": "Node.js",
    "node.js": "Node.js",
    "nodejs": "Node.js",
    "nestjs": "NestJS",
    "spring": "Spring",
    "spring boot": "Spring Boot",
    ".net": ".NET",
    "dotnet": ".NET",
    "rails": "Rails",
    "ruby on rails": "Rails",
    "laravel": "Laravel",
    "graphql": "GraphQL",
    "rest": "REST",
    "rest api": "REST API",
    "restful": "REST",
    # Data / ML
    "pandas": "Pandas",
    "numpy": "NumPy",
    "scikit-learn": "scikit-learn",
    "sklearn": "scikit-learn",
    "tensorflow": "TensorFlow",
    "pytorch": "PyTorch",
    "keras": "Keras",
    "spark": "Spark",
    "hadoop": "Hadoop",
    "airflow": "Airflow",
    "dbt": "dbt",
    "etl": "ETL",
    "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "nlp": "NLP",
    "computer vision": "Computer Vision",
    "data analysis": "Data Analysis",
    "data science": "Data Science",
    "data engineering": "Data Engineering",
    # Cloud / devops
    "aws": "AWS",
    "azure": "Azure",
    "gcp": "GCP",
    "google cloud": "GCP",
    "docker": "Docker",
    "kubernetes": "Kubernetes",
    "k8s": "Kubernetes",
    "terraform": "Terraform",
    "ansible": "Ansible",
    "jenkins": "Jenkins",
    "ci/cd": "CI/CD",
    "cicd": "CI/CD",
    "github actions": "GitHub Actions",
    "gitlab": "GitLab",
    "linux": "Linux",
    "nginx": "Nginx",
    # Databases
    "postgresql": "PostgreSQL",
    "postgres": "PostgreSQL",
    "mysql": "MySQL",
    "mongodb": "MongoDB",
    "redis": "Redis",
    "elasticsearch": "Elasticsearch",
    "dynamodb": "DynamoDB",
    "sqlite": "SQLite",
    "cassandra": "Cassandra",
    "snowflake": "Snowflake",
    "bigquery": "BigQuery",
    # Other
    "git": "Git",
    "github": "GitHub",
    "agile": "Agile",
    "scrum": "Scrum",
    "jira": "Jira",
    "kafka": "Kafka",
    "rabbitmq": "RabbitMQ",
    "celery": "Celery",
    "microservices": "Microservices",
    "api design": "API Design",
    "system design": "System Design",
    "testing": "Testing",
    "pytest": "pytest",
    "jest": "Jest",
    "playwright": "Playwright",
    "selenium": "Selenium",
    "figma": "Figma",
    "ui/ux": "UI/UX",
    "product management": "Product Management",
    "project management": "Project Management",
    "backend": "Backend",
    "frontend": "Frontend",
    "full stack": "Full Stack",
    "fullstack": "Full Stack",
}

# More specific patterns first — preserve role flavour (Backend Developer not generic SE)
TITLE_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (
        re.compile(
            r"\b(senior|sr\.?|lead|principal|staff|junior|jr\.?|mid[\s-]?level|mid)\s+"
            r"(backend|back[\s-]?end|frontend|front[\s-]?end|full[\s-]?stack|software|web|mobile|cloud|devops|platform)\s+"
            r"(engineer|developer|dev)\b",
            re.I,
        ),
        "__MATCH__",  # use matched text title-cased
    ),
    (
        re.compile(
            r"\b(backend|back[\s-]?end|frontend|front[\s-]?end|full[\s-]?stack|software|web|mobile|cloud|devops|platform)\s+"
            r"(engineer|developer|dev)\b",
            re.I,
        ),
        "__MATCH__",
    ),
    (re.compile(r"\bbackend\s+dev(?:eloper)?\b", re.I), "Backend Developer"),
    (re.compile(r"\bfrontend\s+dev(?:eloper)?\b", re.I), "Frontend Developer"),
    (re.compile(r"\bfull[\s-]?stack\s+dev(?:eloper)?\b", re.I), "Full Stack Developer"),
    (re.compile(r"\bsoftware\s+dev(?:eloper)?\b", re.I), "Software Developer"),
    (re.compile(r"\bweb\s+dev(?:eloper)?\b", re.I), "Web Developer"),
    (re.compile(r"\b(data\s+engineer)\b", re.I), "Data Engineer"),
    (re.compile(r"\b(data\s+scientist)\b", re.I), "Data Scientist"),
    (re.compile(r"\b(data\s+analyst)\b", re.I), "Data Analyst"),
    (re.compile(r"\b(machine\s+learning\s+engineer|ml\s+engineer)\b", re.I), "ML Engineer"),
    (re.compile(r"\b(devops\s+engineer)\b", re.I), "DevOps Engineer"),
    (re.compile(r"\b(sre|site\s+reliability\s+engineer)\b", re.I), "SRE"),
    (re.compile(r"\b(cloud\s+engineer)\b", re.I), "Cloud Engineer"),
    (re.compile(r"\b(product\s+manager)\b", re.I), "Product Manager"),
    (re.compile(r"\b(project\s+manager)\b", re.I), "Project Manager"),
    (re.compile(r"\b(qa\s+engineer|quality\s+assurance)\b", re.I), "QA Engineer"),
    (
        re.compile(
            r"\b(mobile\s+(engineer|developer)|ios\s+developer|android\s+developer)\b", re.I
        ),
        "Mobile Developer",
    ),
    (re.compile(r"\b(engineering\s+manager)\b", re.I), "Engineering Manager"),
    (re.compile(r"\b(tech(?:nical)?\s+lead|team\s+lead)\b", re.I), "Tech Lead"),
    (re.compile(r"\b(software\s+engineer)\b", re.I), "Software Engineer"),
]

# Loose headline-ish roles often alone on a line
HEADLINE_LINE_RE = re.compile(
    r"(?im)^\s*((?:senior|sr\.?|lead|junior|jr\.?|mid[\s-]?level)?\s*"
    r"(?:backend|back[\s-]?end|frontend|front[\s-]?end|full[\s-]?stack|software|web|python|java|"
    r"data|devops|cloud|mobile|platform)?\s*"
    r"(?:engineer|developer|dev|programmer|architect|analyst|scientist|manager|lead))\s*$"
)

SECTION_HEADINGS = {
    "experience": r"experience|work\s+experience|employment|work\s+history|professional\s+experience|career",
    "education": r"education|academic|qualifications",
    "skills": r"skills|technical\s+skills|tech\s+stack|technologies|core\s+competencies|competencies",
    "summary": r"summary|profile|about\s+me|objective|professional\s+summary",
}


@dataclass
class ParsedCV:
    raw_text: str
    skills: list[str] = field(default_factory=list)
    titles: list[str] = field(default_factory=list)
    keywords: list[str] = field(default_factory=list)
    email: str | None = None
    phone: str | None = None
    suggested_headline: str | None = None
    experience_text: str | None = None
    education_text: str | None = None
    summary_text: str | None = None
    full_name: str | None = None


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """Extract plain text from PDF, DOCX, or plain text uploads."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith((".txt", ".md", ".rtf")):
        return content.decode("utf-8", errors="ignore")
    try:
        text = _extract_pdf(content)
        if text.strip():
            return text
    except Exception:  # noqa: BLE001
        pass
    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required for PDF upload. pip install pypdf") from exc

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Could not extract text from PDF (may be scanned/image-only)")
    return text


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX upload. pip install python-docx") from exc

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Could not extract text from DOCX")
    return text


def parse_cv_text(text: str) -> ParsedCV:
    """Heuristic parse: skills, titles, sections, contact."""
    if not text or not text.strip():
        return ParsedCV(raw_text=text or "")

    # Normalize weird PDF spacing (B a c k e n d  D e v)
    cleaned = _fix_spaced_letters(text)
    cleaned = cleaned.replace("\u00a0", " ")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    normalized = _normalize(cleaned)
    skills = _extract_skills(normalized, cleaned)
    # Prefer roles near the top of the CV (headline) over ones buried in history
    header_titles = _extract_titles_from_header(cleaned)
    body_titles = _extract_titles(cleaned)
    titles = _merge_unique(header_titles, body_titles)

    email = _extract_email(cleaned)
    phone = _extract_phone(cleaned)
    keywords = _extra_keywords(normalized, skills)
    experience = _section_after(cleaned, SECTION_HEADINGS["experience"], max_chars=2500)
    education = _section_after(cleaned, SECTION_HEADINGS["education"], max_chars=1200)
    summary = _section_after(cleaned, SECTION_HEADINGS["summary"], max_chars=800)
    full_name = _guess_name(cleaned)

    headline = titles[0] if titles else None
    if not headline and skills:
        # soft fallback
        if any(s.lower() in ("backend", "fastapi", "django", "flask", "node.js") for s in skills):
            headline = "Backend Developer"
        elif any(s.lower() in ("react", "vue", "angular", "frontend") for s in skills):
            headline = "Frontend Developer"

    return ParsedCV(
        raw_text=text,
        skills=skills,
        titles=titles,
        keywords=keywords,
        email=email,
        phone=phone,
        suggested_headline=headline,
        experience_text=experience.strip() or None,
        education_text=education.strip() or None,
        summary_text=summary.strip() or None,
        full_name=full_name,
    )


def cv_profile_text(cv) -> str:
    """Combine structured CV fields + raw upload text for parsing/matching."""
    parts = [
        cv.full_name,
        cv.headline,
        cv.summary,
        cv.experience,
        cv.education,
        cv.skills,
        cv.links,
        cv.raw_text,
        cv.parsed_skills,
        cv.parsed_titles,
    ]
    return "\n".join(p for p in parts if p)


def _fix_spaced_letters(text: str) -> str:
    """Collapse 'B a c k e n d' style PDF extractions into 'Backend'."""

    def repl(match: re.Match[str]) -> str:
        return match.group(0).replace(" ", "")

    # sequences of single letters separated by spaces (at least 3 letters)
    return re.sub(r"\b(?:[A-Za-z]\s+){2,}[A-Za-z]\b", repl, text)


def _normalize(text: str) -> str:
    t = text.lower()
    t = t.replace("\u00a0", " ")
    t = re.sub(r"[/|•·,;()\[\]{}]", " ", t)
    t = re.sub(r"\s+", " ", t)
    return t


def _title_case_role(raw: str) -> str:
    raw = re.sub(r"\s+", " ", raw.strip())
    # normalize back-end → Backend, full-stack → Full Stack
    raw = re.sub(r"back[\s-]?end", "Backend", raw, flags=re.I)
    raw = re.sub(r"front[\s-]?end", "Frontend", raw, flags=re.I)
    raw = re.sub(r"full[\s-]?stack", "Full Stack", raw, flags=re.I)
    raw = re.sub(r"\bdev\b", "Developer", raw, flags=re.I)
    raw = re.sub(r"\bsr\.?\b", "Senior", raw, flags=re.I)
    raw = re.sub(r"\bjr\.?\b", "Junior", raw, flags=re.I)
    parts = []
    for w in raw.split():
        if w.isupper() and len(w) <= 3:
            parts.append(w)
        else:
            parts.append(w[:1].upper() + w[1:].lower() if w else w)
    return " ".join(parts)


def _extract_skills(normalized: str, original: str) -> list[str]:
    found: dict[str, str] = {}

    items = sorted(SKILL_CATALOG.items(), key=lambda kv: len(kv[0]), reverse=True)
    for key, label in items:
        if " " in key or "." in key or "/" in key or "+" in key or "#" in key:
            if key in normalized:
                found[label] = label
        else:
            if re.search(rf"(?<![a-z0-9]){re.escape(key)}(?![a-z0-9])", normalized):
                found[label] = label

    section = _section_after(original, SECTION_HEADINGS["skills"], max_chars=1200)
    if section:
        for chunk in re.split(r"[,|/•\n;·]+", section):
            token = chunk.strip().lower()
            token = re.sub(r"\s+", " ", token)
            if not token or len(token) > 40:
                continue
            if token in SKILL_CATALOG:
                found[SKILL_CATALOG[token]] = SKILL_CATALOG[token]
            else:
                # multi-word partial: "python django" → scan each
                for part in token.split():
                    if part in SKILL_CATALOG:
                        found[SKILL_CATALOG[part]] = SKILL_CATALOG[part]

    # Prefer concrete tools over generic Backend/Frontend labels if both present
    return list(found.values())


def _extract_titles(text: str) -> list[str]:
    found: list[str] = []
    seen: set[str] = set()
    for pattern, label in TITLE_PATTERNS:
        for m in pattern.finditer(text):
            if label == "__MATCH__":
                title = _title_case_role(m.group(0))
            else:
                title = label
            key = title.lower()
            if key not in seen:
                found.append(title)
                seen.add(key)
    return found


def _extract_titles_from_header(text: str) -> list[str]:
    """Look at early lines for role headlines (common CV layout)."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()][:15]
    found: list[str] = []
    seen: set[str] = set()

    def add(title: str) -> None:
        key = title.lower()
        if key not in seen:
            found.append(title)
            seen.add(key)

    for ln in lines:
        # skip contact lines
        if "@" in ln or (re.search(r"\d{3,}", ln) and len(ln) < 40):
            continue
        if len(ln) > 80:
            continue
        m = HEADLINE_LINE_RE.match(ln)
        if m:
            add(_title_case_role(m.group(1)))
            continue
        # free-form: line contains backend/developer etc.
        if re.search(
            r"\b(backend|frontend|full[\s-]?stack|software)\s+(developer|engineer|dev)\b",
            ln,
            re.I,
        ):
            m2 = re.search(
                r"((?:senior|sr\.?|junior|jr\.?|lead)?\s*"
                r"(?:backend|frontend|full[\s-]?stack|software)\s+"
                r"(?:developer|engineer|dev))",
                ln,
                re.I,
            )
            if m2:
                add(_title_case_role(m2.group(1)))
    return found


def _guess_name(text: str) -> str | None:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()][:8]
    for ln in lines:
        if "@" in ln or re.search(r"https?://", ln, re.I):
            continue
        if re.search(r"\d{5,}", ln):
            continue
        # 2–4 capitalized words
        if re.match(r"^[A-Z][a-zA-Z'’\-]+(?:\s+[A-Z][a-zA-Z'’\-]+){1,3}$", ln):
            if not re.search(
                r"engineer|developer|manager|summary|experience|education|skills",
                ln,
                re.I,
            ):
                return ln
    return None


def _extract_email(text: str) -> str | None:
    m = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I)
    return m.group(0) if m else None


def _extract_phone(text: str) -> str | None:
    m = re.search(
        r"(?:\+\d{1,3}[\s-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}",
        text,
    )
    if not m:
        return None
    candidate = m.group(0).strip()
    digits = re.sub(r"\D", "", candidate)
    if len(digits) < 7:
        return None
    return candidate


def _extra_keywords(normalized: str, skills: list[str]) -> list[str]:
    extras = []
    for phrase in (
        "remote",
        "backend",
        "frontend",
        "full stack",
        "fullstack",
        "api",
        "microservices",
        "distributed systems",
        "cloud",
        "security",
        "fintech",
        "healthcare",
        "saas",
    ):
        if phrase in normalized:
            extras.append(phrase.title() if phrase != "api" else "API")
    skill_set = {s.lower() for s in skills}
    return [e for e in extras if e.lower() not in skill_set][:12]


def _section_after(text: str, heading_re: str, max_chars: int = 800) -> str:
    """Grab text after a section heading until the next known heading."""
    m = re.search(rf"(?im)^(?:{heading_re})\s*:?\s*$", text)
    if not m:
        m = re.search(rf"(?im)^(?:{heading_re})\s*[:\-–—]\s*", text)
    if not m:
        m = re.search(rf"(?i)(?:{heading_re})\s*[:\-–—]\s*", text)
    if not m:
        return ""
    start = m.end()
    rest = text[start : start + max_chars]
    # next section heading
    next_heads = "|".join(SECTION_HEADINGS.values())
    stop = re.search(rf"(?im)^(?:{next_heads})\s*:?\s*$", rest)
    if stop and stop.start() > 10:
        return rest[: stop.start()].strip()
    return rest.strip()


def _merge_unique(primary: list[str], extra: list[str]) -> list[str]:
    seen = {p.lower() for p in primary}
    out = list(primary)
    for e in extra:
        if e.lower() not in seen:
            out.append(e)
            seen.add(e.lower())
    return out
